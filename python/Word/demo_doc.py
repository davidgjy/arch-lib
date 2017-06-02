import docx

doc = docx.Document('demo.docx')
print('paragraphs number: %s' % len(doc.paragraphs))
print('1st paragraph: %s' % doc.paragraphs[0].text)
print('2nd paragraph: %s' % doc.paragraphs[1].text)
print('paragraphs runs: %s' % len(doc.paragraphs[1].runs))
print('1st paragraph run: %s' % doc.paragraphs[1].runs[0].text)
print('2nd paragraph run: %s' % doc.paragraphs[1].runs[1].text)
print('3rd paragraph run: %s' % doc.paragraphs[1].runs[2].text)
print('4th paragraph run: %s' % doc.paragraphs[1].runs[3].text)
